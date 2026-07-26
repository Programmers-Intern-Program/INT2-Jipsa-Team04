"""DOCX 본문의 문단, 제목, 목록과 표를 원본 블록 순서대로 추출한다.

DOCX는 단순한 텍스트 파일이 아니라 여러 XML 파트와 관계 파일을 포함하는
OOXML ZIP 패키지다. 이 파서는 ``python-docx``를 이용하여 본문 XML의 직계
자식 블록을 순회하고, 문단과 표가 실제 문서에 나타나는 순서를 보존한다.

추출 범위는 다음과 같다.

- 일반 문단 텍스트
- 기본 또는 사용자 정의 스타일을 통해 상속된 제목 레벨
- 문단 XML 또는 스타일 XML에 연결된 글머리 기호·번호 목록
- 표의 행·열 텍스트
- 섹션 전환 위치
- 본문 전체의 문단, 표, 블록 개수

머리글, 바닥글, 각주, 미주, 주석, 텍스트 상자와 이미지 OCR은 현재 범위에
포함하지 않는다. 해당 요소를 지원할 때는 원본 위치 메타데이터 계약과 파서
버전을 함께 변경해야 한다.
"""

import asyncio
import re
from collections.abc import Mapping
from pathlib import Path
from typing import Final, cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

from jipsa_rag.infrastructure.document.exceptions import (
    DocumentParserError,
    DocumentReadError,
    DocumentSourceMetadataValue,
    DocumentTextExtractionError,
    DocumentTextNotFoundError,
    InvalidDocumentError,
)
from jipsa_rag.infrastructure.document.models import (
    DocumentType,
    ParsedDocument,
    ParsedDocumentUnit,
    SourceMetadataValue,
)
from jipsa_rag.infrastructure.document.parsers._common import (
    normalize_inline_text,
    normalize_text,
    validate_ooxml_package,
)

# Local RAG DB의 Parser_Type에 저장할 안정적인 식별자다. 단순히 DOCX라고만
# 기록하지 않고 텍스트 기반 추출 방식임을 명시한다.
_DOCX_PARSER_TYPE: Final[str] = "DOCX_TEXT"

# 텍스트 정규화, 블록 경계 또는 source_metadata 의미가 바뀌면 증가시킨다.
# 이번 변경은 주석 보강만 포함하므로 기존 결과 호환 버전 1.0.0을 유지한다.
_DOCX_PARSER_VERSION: Final[str] = "1.0.0"

# 유효한 DOCX 패키지라면 공통 Content Types와 본문 루트가 반드시 존재해야 한다.
# ZIP Magic Byte만으로는 PPTX/XLSX와 구분할 수 없으므로 내부 경로까지 확인한다.
_DOCX_REQUIRED_MEMBERS: Final[frozenset[str]] = frozenset(
    {
        "[Content_Types].xml",
        "word/document.xml",
    }
)

# Word 기본 영문 스타일명은 Heading 1, 한국어 UI에서는 제목 1 형태가 일반적이다.
# 사용자 정의 스타일은 기본 Heading 스타일을 상속할 수 있으므로 실제 탐색은
# 현재 스타일에서 base_style 체인을 따라가며 수행한다.
_HEADING_STYLE_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"^(?:heading|제목)\s*([1-9]\d*)$",
    re.IGNORECASE,
)


def _find_numbering_properties(
    paragraph_properties: BaseOxmlElement | None,
) -> BaseOxmlElement | None:
    """문단 속성에서 ``w:numPr`` 자식을 타입 안전하게 반환한다.

    python-docx의 ``numPr`` descriptor는 런타임에서는 XML 요소를 반환하지만
    정적 타입 정의에서는 ``ZeroOrOne`` descriptor로 노출될 수 있다. 이 함수는
    OOXML QName 기반 ``find()``를 사용하여 descriptor 타입에 의존하지 않는다.
    """

    if paragraph_properties is None:
        return None

    return cast(
        BaseOxmlElement | None,
        paragraph_properties.find(qn("w:numPr")),
    )


def _read_numbering_integer(
    numbering_properties: BaseOxmlElement,
    *,
    child_name: str,
) -> int | None:
    """``w:numPr`` 아래 정수 속성을 읽고 누락 값은 ``None``으로 반환한다.

    ``numId``와 ``ilvl``은 모두 자식 요소의 ``w:val`` 속성에 저장된다. 손상된
    문서에서 요소나 속성이 누락되면 호출자가 목록 없음 또는 기본 레벨로 처리할
    수 있도록 예외 대신 ``None``을 반환한다. 숫자가 아닌 값은 상위 파서 경계에서
    손상 문서 예외로 변환되도록 ``int()``의 ``ValueError``를 그대로 유지한다.
    """

    child = numbering_properties.find(qn(child_name))
    if child is None:
        return None

    raw_value = child.get(qn("w:val"))
    if raw_value is None:
        return None

    return int(raw_value)


class DocxDocumentParser:
    """``python-docx``를 사용하여 DOCX 본문 블록과 위치 정보를 추출한다.

    하나의 문단 또는 하나의 표를 하나의 ``ParsedDocumentUnit``으로 만든다.
    이 경계를 유지하면 후속 ``CharacterTextChunker``가 서로 다른 문단이나 표를
    하나의 청크에 강제로 합치지 않고, 각 청크에 해당 원본 위치 메타데이터를
    그대로 전달할 수 있다.
    """

    @property
    def file_type(self) -> DocumentType:
        """이 구현체가 처리하는 공통 문서 형식인 DOCX를 반환한다."""

        return DocumentType.DOCX

    @property
    def parser_type(self) -> str:
        """Local RAG DB에 저장할 DOCX 텍스트 파서 종류를 반환한다."""

        return _DOCX_PARSER_TYPE

    @property
    def parser_version(self) -> str:
        """재파싱 및 결정적 Chunk ID 생성에 사용할 파서 버전을 반환한다."""

        return _DOCX_PARSER_VERSION

    async def parse(self, file_path: Path) -> ParsedDocument:
        """동기식 DOCX 파싱을 작업 스레드로 이동하여 비동기로 제공한다.

        ``python-docx``는 파일 읽기와 XML 파싱을 동기 방식으로 수행한다. 이를
        FastAPI 이벤트 루프에서 직접 실행하면 큰 문서 처리 중 다른 요청이 같은
        이벤트 루프를 사용하지 못할 수 있으므로 ``asyncio.to_thread()``를 사용한다.

        Args:
            file_path:
                다운로드와 기본 파일 형식 검증이 완료된 임시 DOCX 경로다.

        Returns:
            본문 블록별 텍스트와 DOCX 위치 메타데이터를 포함한 결과다.
        """

        return await asyncio.to_thread(self._parse_sync, file_path)

    def _parse_sync(self, file_path: Path) -> ParsedDocument:
        """DOCX 패키지를 열고 본문 블록을 공통 결과 모델로 변환한다."""

        # 파서가 독립적으로 호출되는 상황도 고려하여, 다운로드 계층의 검증 여부와
        # 관계없이 자신의 경계에서 DOCX 패키지 루트를 다시 확인한다.
        validate_ooxml_package(
            file_path,
            file_type=self.file_type,
            required_members=_DOCX_REQUIRED_MEMBERS,
        )

        try:
            # Document()가 반환하는 객체는 열린 파일 경로에 계속 의존하지 않도록
            # _parse_body() 실행 안에서 필요한 문자열과 메타데이터로 모두 변환한다.
            document = Document(str(file_path))

            # 목록 유형을 정확히 구분하려면 본문 문단의 numId뿐 아니라
            # numbering.xml의 abstractNum/level별 numFmt를 함께 해석해야 한다.
            numbering_formats = self._build_numbering_format_map(document)

            units, paragraph_count, table_count = self._parse_body(
                document,
                numbering_formats=numbering_formats,
            )
        except DocumentParserError:
            # 이미 공통 문서 예외로 분류된 오류는 의미를 잃지 않도록 그대로 전달한다.
            raise
        except OSError as error:
            # 권한, 파일 잠금, 장치 오류처럼 파일 시스템 계층에서 발생한 문제다.
            raise DocumentReadError(file_path) from error
        except Exception as error:
            # python-docx와 lxml은 손상된 관계, XML 또는 패키지에 대해 여러 내부
            # 예외를 발생시킬 수 있다. 라이브러리 세부 타입을 API에 노출하지 않고
            # 공통 손상 문서 예외로 통일한다.
            raise InvalidDocumentError(self.file_type) from error

        parsed_document = ParsedDocument(
            file_type=self.file_type,
            units=tuple(units),
            document_metadata={
                # python-docx가 계산한 섹션 수다. 각 unit의 section_index와 함께
                # 문서 전체 및 원본 블록 단위 위치를 모두 확인할 수 있다.
                "section_count": len(document.sections),
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "block_count": len(units),
            },
        )

        # 빈 문단이나 빈 표는 원본 위치 보존을 위해 unit으로 남을 수 있다. 그러나
        # 문서 전체에서 실제 검색 가능한 텍스트가 하나도 없다면 인제스트를 성공으로
        # 처리하지 않고 명시적인 텍스트 없음 오류를 반환한다.
        if parsed_document.text_unit_count == 0:
            raise DocumentTextNotFoundError(self.file_type)

        return parsed_document

    def _parse_body(
        self,
        document: DocxDocument,
        *,
        numbering_formats: dict[tuple[int, int], str],
    ) -> tuple[list[ParsedDocumentUnit], int, int]:
        """본문 XML의 문단과 표를 실제 등장 순서대로 순회한다.

        ``document.paragraphs``와 ``document.tables``를 각각 순회하면 두 종류의
        상대적인 원본 순서를 잃는다. 따라서 ``word/document.xml`` 본문의 직계
        자식 요소를 한 번만 순회하고 ``CT_P``와 ``CT_Tbl``을 그 자리에서 처리한다.

        Args:
            document:
                ``python-docx``가 연 DOCX 문서 객체다.
            numbering_formats:
                ``(numId, list_level)``을 Word 목록 형식 문자열로 매핑한 사전이다.

        Returns:
            파싱 unit 목록, 전체 문단 수와 전체 표 수를 반환한다.
        """

        units: list[ParsedDocumentUnit] = []

        # 사용자에게 표시할 위치 값은 모두 1부터 시작한다. section_index는 현재
        # 블록이 속한 섹션을 나타내고, block_index는 문단과 표를 합친 본문 순번이다.
        section_index = 1
        paragraph_index = 0
        table_index = 0

        for block_index, child in enumerate(
            document.element.body.iterchildren(),
            start=1,
        ):
            # 오류 예외에는 JSON scalar만 허용된다. 이 기본 위치 정보는 정수만
            # 저장하므로 DocumentTextExtractionError의 메타데이터 계약을 정확히
            # 만족하도록 더 좁은 값 타입으로 선언한다.
            base_metadata: dict[str, DocumentSourceMetadataValue] = {
                "section_index": section_index,
                "block_index": block_index,
            }

            try:
                if isinstance(child, CT_P):
                    paragraph_index += 1

                    # Paragraph 프록시에 document를 parent로 전달해야 스타일, 관계와
                    # 번호 매기기 속성을 정상적으로 조회할 수 있다.
                    paragraph = Paragraph(child, document)

                    units.append(
                        self._parse_paragraph(
                            paragraph,
                            paragraph_index=paragraph_index,
                            numbering_formats=numbering_formats,
                            base_metadata=base_metadata,
                        )
                    )

                    # Word에서는 문단 속성 안의 sectPr가 "이 문단 뒤에서 현재 섹션이
                    # 끝난다"는 의미다. 따라서 현재 문단에는 기존 section_index를
                    # 기록하고, 다음 본문 블록부터 번호를 증가시킨다.
                    paragraph_properties = paragraph._p.pPr
                    if paragraph_properties is not None and paragraph_properties.sectPr is not None:
                        section_index += 1

                elif isinstance(child, CT_Tbl):
                    table_index += 1
                    table = Table(child, document)

                    units.append(
                        self._parse_table(
                            table,
                            table_index=table_index,
                            base_metadata=base_metadata,
                        )
                    )

                # sectPr처럼 문단과 표가 아닌 본문 자식은 텍스트 unit으로 만들지 않는다.
                # block_index는 원본 XML 위치를 반영하므로 해당 요소를 포함한 순번을 유지한다.

            except DocumentParserError:
                raise
            except Exception as error:
                # 특정 본문 블록에서만 추출이 실패한 경우 문서 전체 손상으로 뭉개지 않고
                # section_index와 block_index를 포함한 위치별 추출 오류로 변환한다.
                raise DocumentTextExtractionError(
                    file_type=self.file_type,
                    source_metadata=base_metadata,
                ) from error

        return units, paragraph_index, table_index

    def _parse_paragraph(
        self,
        paragraph: Paragraph,
        *,
        paragraph_index: int,
        numbering_formats: dict[tuple[int, int], str],
        base_metadata: Mapping[str, SourceMetadataValue],
    ) -> ParsedDocumentUnit:
        """단일 문단 텍스트와 제목·목록 분류 메타데이터를 생성한다."""

        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading_level = self._extract_heading_level(paragraph)
        list_info = self._extract_list_info(paragraph, numbering_formats)

        metadata: dict[str, SourceMetadataValue] = {
            **base_metadata,
            "paragraph_index": paragraph_index,
            # 스타일이 없는 문단은 빈 문자열보다 명시적인 None으로 기록한다.
            "style_name": style_name or None,
        }

        if heading_level is not None:
            # 제목 스타일과 목록 속성이 동시에 존재하는 비정상 문서에서는 제목을
            # 우선한다. 제목 레벨은 문서 계층과 검색 결과 표시에서 더 중요한 구조다.
            metadata.update(
                {
                    "unit_type": "heading",
                    "heading_level": heading_level,
                }
            )
        elif list_info is not None:
            num_id, list_level, list_format = list_info
            metadata.update(
                {
                    "unit_type": "list_item",
                    "list_numbering_id": num_id,
                    # Word의 ilvl은 0부터 시작하는 내부 목록 깊이다. 원본 XML 값과
                    # 일치하도록 그대로 보존한다.
                    "list_level": list_level,
                    "list_type": ("bullet" if list_format == "bullet" else "numbered"),
                    # decimal, lowerLetter, upperRoman 등 실제 numFmt 값도 함께 남긴다.
                    "list_format": list_format,
                }
            )
        else:
            metadata["unit_type"] = "paragraph"

        return ParsedDocumentUnit(
            text=normalize_text(paragraph.text),
            source_metadata=metadata,
        )

    @staticmethod
    def _parse_table(
        table: Table,
        *,
        table_index: int,
        base_metadata: Mapping[str, SourceMetadataValue],
    ) -> ParsedDocumentUnit:
        """DOCX 표를 탭 구분 행 문자열과 위치 메타데이터로 변환한다.

        표의 각 행은 줄바꿈으로, 각 셀은 탭으로 연결한다. 셀 내부의 여러 문단은
        ``normalize_inline_text()``가 공백 하나로 연결하므로 표의 행 경계가 유지된다.
        """

        rows: list[str] = []
        maximum_column_count = 0

        for row in table.rows:
            # python-docx는 병합 셀을 시각적 열 위치에 맞춰 반복해서 반환할 수 있다.
            # 현재 계약은 렌더링된 행 텍스트를 보존하는 것이므로 row.cells 순서를 따른다.
            cells = [normalize_inline_text(cell.text) for cell in row.cells]
            maximum_column_count = max(maximum_column_count, len(cells))

            # 오른쪽의 완전히 빈 셀 때문에 불필요한 탭이 남지 않도록 rstrip()을
            # 적용한다. 중간 빈 셀의 탭은 열 위치를 나타내므로 유지한다.
            rows.append("\t".join(cells).rstrip())

        return ParsedDocumentUnit(
            text=normalize_text("\n".join(rows)),
            source_metadata={
                **base_metadata,
                "unit_type": "table",
                "table_index": table_index,
                "row_count": len(table.rows),
                "column_count": maximum_column_count,
            },
        )

    @staticmethod
    def _extract_heading_level(paragraph: Paragraph) -> int | None:
        """현재 문단 스타일과 상속 스타일에서 제목 레벨을 찾는다.

        사용자 정의 스타일은 직접 ``Heading 1``이라는 이름을 사용하지 않아도
        기본 제목 스타일을 ``base_style``로 상속할 수 있다. 현재 스타일 하나만
        확인하면 이러한 문단을 일반 문단으로 오인하므로 상속 체인을 끝까지 따라간다.
        순환 참조가 있는 손상 스타일을 방어하기 위해 방문한 style_id를 기록한다.
        """

        style = paragraph.style
        visited_style_ids: set[str] = set()

        while style is not None and style.style_id not in visited_style_ids:
            visited_style_ids.add(style.style_id)

            # 지역화된 표시 이름과 내부 style_id를 모두 확인한다. style_id의 밑줄은
            # 공백으로 바꿔 "Heading_1" 형태도 같은 규칙으로 처리한다.
            candidates = (
                style.name.strip(),
                style.style_id.replace("_", " ").strip(),
            )

            for candidate in candidates:
                match = _HEADING_STYLE_PATTERN.match(candidate)
                if match is not None:
                    return int(match.group(1))

            style = style.base_style

        return None

    @staticmethod
    def _extract_list_info(
        paragraph: Paragraph,
        numbering_formats: dict[tuple[int, int], str],
    ) -> tuple[int, int, str] | None:
        """문단에 적용된 목록 번호 ID, 깊이와 형식을 반환한다.

        Word의 목록 연결은 두 위치 중 하나에 존재할 수 있다.

        1. 문단 자체의 ``w:pPr/w:numPr``
        2. ``List Bullet`` 같은 적용 스타일 또는 상속 스타일의 ``w:pPr/w:numPr``

        문단 직접 속성을 우선하고, 없을 때만 스타일 상속 체인을 탐색한다. 이 순서는
        사용자가 특정 문단에서 스타일의 목록 설정을 덮어쓴 경우를 정확히 반영한다.
        """

        paragraph_properties = paragraph._p.pPr
        numbering_properties = _find_numbering_properties(paragraph_properties)

        if numbering_properties is None:
            style = paragraph.style
            visited_style_ids: set[str] = set()

            while style is not None and style.style_id not in visited_style_ids:
                visited_style_ids.add(style.style_id)

                # python-docx의 xmlchemy descriptor 타입은 Mypy에서 실제 XML
                # 자식 속성(numId, ilvl)을 충분히 표현하지 못한다. descriptor에
                # 직접 접근하지 않고 표준 OOXML 자식 경로를 조회하면 런타임 동작과
                # 정적 타입 계약을 동시에 유지할 수 있다.
                numbering_properties = _find_numbering_properties(
                    style.element.pPr,
                )

                if numbering_properties is not None:
                    break

                style = style.base_style

        if numbering_properties is None:
            return None

        num_id = _read_numbering_integer(
            numbering_properties,
            child_name="w:numId",
        )

        # Word는 numId=0을 목록 번호 제거 또는 목록 없음 값으로 사용할 수 있다.
        if num_id is None or num_id == 0:
            return None

        list_level = (
            _read_numbering_integer(
                numbering_properties,
                child_name="w:ilvl",
            )
            or 0
        )

        # numbering.xml에 구체 형식이 없거나 손상된 경우에도 목록 위치는 보존한다.
        # 기본값 decimal을 사용하되, bullet이 확인되면 list_type을 bullet로 분류한다.
        list_format = numbering_formats.get((num_id, list_level), "decimal")

        return num_id, list_level, list_format

    @staticmethod
    def _build_numbering_format_map(
        document: DocxDocument,
    ) -> dict[tuple[int, int], str]:
        """numbering.xml을 해석하여 ``(numId, level)``별 형식을 계산한다.

        Word 번호 매기기는 다음과 같은 간접 참조 구조를 사용한다.

        ``w:num/@numId`` → ``w:abstractNumId/@w:val`` →
        ``w:abstractNum/w:lvl/@w:ilvl`` → ``w:numFmt/@w:val``

        따라서 본문 문단에 있는 numId만으로는 글머리 기호인지 decimal 번호인지
        알 수 없다. 두 단계의 임시 맵을 만든 뒤 실제 numId와 level 조합으로 합친다.
        """

        try:
            numbering_root = document.part.numbering_part.element
        except (KeyError, AttributeError):
            # 번호 매기기 파트가 없는 문서는 정상적인 일반 문서일 수 있다.
            return {}

        # 실제 목록 인스턴스 numId가 어떤 abstractNum 정의를 참조하는지 저장한다.
        abstract_id_by_num_id: dict[int, int] = {}

        # abstractNum 정의 안에서 각 level이 사용하는 numFmt를 저장한다.
        format_by_abstract_level: dict[tuple[int, int], str] = {}

        for num_element in numbering_root.findall(qn("w:num")):
            num_id_value = num_element.get(qn("w:numId"))
            abstract_reference = num_element.find(qn("w:abstractNumId"))

            if num_id_value is None or abstract_reference is None:
                continue

            abstract_id_value = abstract_reference.get(qn("w:val"))
            if abstract_id_value is not None:
                abstract_id_by_num_id[int(num_id_value)] = int(abstract_id_value)

        for abstract_element in numbering_root.findall(qn("w:abstractNum")):
            abstract_id_value = abstract_element.get(qn("w:abstractNumId"))
            if abstract_id_value is None:
                continue

            abstract_id = int(abstract_id_value)

            for level_element in abstract_element.findall(qn("w:lvl")):
                level_value = level_element.get(qn("w:ilvl"))
                number_format = level_element.find(qn("w:numFmt"))

                if level_value is None or number_format is None:
                    continue

                format_value = number_format.get(qn("w:val"))
                if format_value is not None:
                    format_by_abstract_level[(abstract_id, int(level_value))] = format_value

        # 실제 numId마다 참조하는 abstractNum의 모든 level 형식을 펼쳐 최종 맵을 만든다.
        return {
            (num_id, level): number_format
            for num_id, abstract_id in abstract_id_by_num_id.items()
            for (
                candidate_abstract_id,
                level,
            ), number_format in format_by_abstract_level.items()
            if candidate_abstract_id == abstract_id
        }
